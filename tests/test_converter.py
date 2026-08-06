import unittest
import os
import csv
import shutil
import docx
from docx_md_converter.slugify import normalize_slug
from docx_md_converter.parser import parse_docx_file
from docx_md_converter.writer import write_markdown_file, build_filename
from docx_md_converter.converter import convert_single_docx, convert_batch
from docx_md_converter.history import ConversionHistory

class TestDocxMdConverter(unittest.TestCase):

    def setUp(self):
        self.test_dir = "tests/tmp_test_files"
        self.output_dir = "tests/tmp_output_md"
        self.csv_history = "tests/tmp_conversion_history.csv"
        os.makedirs(self.test_dir, exist_ok=True)
        os.makedirs(self.output_dir, exist_ok=True)

    def tearDown(self):
        if os.path.exists(self.test_dir):
            shutil.rmtree(self.test_dir)
        if os.path.exists(self.output_dir):
            shutil.rmtree(self.output_dir)
        if os.path.exists(self.csv_history):
            os.remove(self.csv_history)

    def test_normalize_slug(self):
        self.assertEqual(normalize_slug("Família de Produto"), "familia-de-produto")
        self.assertEqual(normalize_slug("Conta Corrente"), "conta-corrente")
        self.assertEqual(normalize_slug("Abertura Digital (PF)"), "abertura-digital-pf")
        self.assertEqual(normalize_slug("Módulo: Teste"), "modulo-teste")
        self.assertEqual(normalize_slug("Palavras chave"), "palavras-chave")

    def test_sample_docx_conversion(self):
        docx_path = os.path.join(self.test_dir, "sample.docx")
        doc = docx.Document()
        
        doc.add_paragraph("Manual de Integração do Módulo")
        doc.add_paragraph("- BU: Varejo Especial")
        doc.add_paragraph("- Família de Produto: Conta Corrente")
        doc.add_paragraph("- Produto: Abertura Digital")
        doc.add_paragraph("- Subproduto: PF")
        doc.add_paragraph("- Módulo: Cadastro Inicial")
        doc.add_paragraph("- Segmento: Premium")
        doc.add_paragraph("- Palavras chave: abertura, conta, pf")
        
        doc.add_heading("Visão Geral do Processo", level=1)
        doc.add_paragraph("Este documento descreve as etapas para abertura de conta.")
        
        doc.save(docx_path)
        
        out_tuple = convert_single_docx(docx_path, output_root=self.output_dir)
        output_file = out_tuple[0]
        
        expected_folder = os.path.join(self.output_dir, "conta-corrente")
        expected_filename = "varejo-especial-conta-corrente-abertura-digital-pf-cadastro-inicial-premium.md"
        expected_filepath = os.path.join(expected_folder, expected_filename)
        
        self.assertTrue(os.path.exists(expected_filepath), f"Arquivo não encontrado: {expected_filepath}")
        
        with open(expected_filepath, "r", encoding="utf-8") as f:
            content = f.read()
            
        self.assertIn('---', content)
        self.assertIn('title: "varejo-especial | conta-corrente | abertura-digital | pf | cadastro-inicial | premium"', content)
        self.assertIn('bu: varejo-especial', content)
        self.assertIn('familia-de-produto: conta-corrente', content)
        self.assertIn('produto: abertura-digital', content)
        self.assertIn('subproduto: pf', content)
        self.assertIn('modulo: cadastro-inicial', content)
        self.assertIn('segmento: premium', content)
        self.assertIn('palavras-chave: abertura-conta-pf', content)
        self.assertIn('# Visão Geral do Processo', content)
        self.assertIn('Este documento descreve as etapas para abertura de conta.', content)

    def test_incremental_conversion_and_csv_history(self):
        # 1. Cria 2 arquivos docx
        doc1_path = os.path.join(self.test_dir, "doc1.docx")
        doc1 = docx.Document()
        doc1.add_paragraph("Documento Um")
        doc1.add_paragraph("- BU: Varejo")
        doc1.add_paragraph("- Família de Produto: Cartões")
        doc1.add_paragraph("- Produto: Credito")
        doc1.add_paragraph("- Subproduto: Visa")
        doc1.add_paragraph("- Módulo: Fatura")
        doc1.save(doc1_path)

        doc2_path = os.path.join(self.test_dir, "doc2.docx")
        doc2 = docx.Document()
        doc2.add_paragraph("Documento Dois")
        doc2.add_paragraph("- BU: Varejo")
        doc2.add_paragraph("- Família de Produto: Cartões")
        doc2.add_paragraph("- Produto: Debito")
        doc2.add_paragraph("- Subproduto: Master")
        doc2.add_paragraph("- Módulo: Extrato")
        doc2.save(doc2_path)

        # Rodada 1: Converte os 2 arquivos
        res1 = convert_batch([self.test_dir], output_root=self.output_dir, history_file=self.csv_history)
        self.assertEqual(len(res1["converted"]), 2)
        self.assertEqual(len(res1["skipped"]), 0)
        self.assertTrue(os.path.exists(self.csv_history))

        # Verifica se as colunas de metadados estão presentes no CSV
        with open(self.csv_history, "r", encoding="utf-8-sig") as f:
            reader = csv.DictReader(f)
            fieldnames = reader.fieldnames
            self.assertIn("bu", fieldnames)
            self.assertIn("familia-de-produto", fieldnames)
            self.assertIn("produto", fieldnames)
            self.assertIn("subproduto", fieldnames)
            self.assertIn("modulo", fieldnames)
            rows = list(reader)
            self.assertEqual(len(rows), 2)
            self.assertEqual(rows[0]["bu"], "varejo")
            self.assertEqual(rows[0]["familia-de-produto"], "cartoes")

        # Rodada 2: Executa novamente sem novos arquivos (deve pular os 2)
        res2 = convert_batch([self.test_dir], output_root=self.output_dir, history_file=self.csv_history)
        self.assertEqual(len(res2["converted"]), 0)
        self.assertEqual(len(res2["skipped"]), 2)

        # Rodada 3: Adiciona doc3
        doc3_path = os.path.join(self.test_dir, "doc3.docx")
        doc3 = docx.Document()
        doc3.add_paragraph("Documento Três")
        doc3.add_paragraph("- BU: Empresas")
        doc3.add_paragraph("- Família de Produto: Empréstimos")
        doc3.add_paragraph("- Produto: Giro")
        doc3.add_paragraph("- Subproduto: PJ")
        doc3.add_paragraph("- Módulo: Simulação")
        doc3.save(doc3_path)

        res3 = convert_batch([self.test_dir], output_root=self.output_dir, history_file=self.csv_history)
        self.assertEqual(len(res3["converted"]), 1)
        self.assertEqual(len(res3["skipped"]), 2)

        # Rodada 4: Usa force=True (deve reprocessar todos os 3)
        res4 = convert_batch([self.test_dir], output_root=self.output_dir, history_file=self.csv_history, force=True)
        self.assertEqual(len(res4["converted"]), 3)

if __name__ == "__main__":
    unittest.main()
