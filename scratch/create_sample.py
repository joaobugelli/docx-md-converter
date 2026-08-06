import docx

doc = docx.Document()
doc.add_paragraph("Especificação Técnica de Abertura de Conta")

doc.add_paragraph("- BU: Varejo")
doc.add_paragraph("- Família de Produto: Conta Corrente")
doc.add_paragraph("- Produto: Abertura Digital")
doc.add_paragraph("- Subproduto: PF")
doc.add_paragraph("- Módulo: Validação Cadastral")
doc.add_paragraph("- Segmento: Uniclass")
doc.add_paragraph("- Palavras chave: conta, corrente, abertura, pf")

doc.add_heading("1. Introdução", level=1)
doc.add_paragraph("Esta seção descreve a validação cadastral para o fluxo de abertura de conta digital.")

doc.add_heading("2. Requisitos", level=2)
doc.add_paragraph("- Validar documento de identidade (RG/CNH)")
doc.add_paragraph("- Validar biometria facial")
doc.add_paragraph("- Validar comprovante de residência")

doc.save("input_docx/exemplo_conta_corrente.docx")
print("Arquivo exemplo_conta_corrente.docx gerado em input_docx/")
